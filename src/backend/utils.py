import re
import os
from typing import List, Dict, Any
import pypdfium2 as pdfium
import docx
import openai
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Set OpenAI API key from environment variable
openai.api_key = os.getenv("OPENAI_API_KEY")

def extract_text(file_path: str) -> str:
    """
    Extract text from PDF or DOCX files.
    
    Args:
        file_path: Path to the document file
        
    Returns:
        Extracted text as a string
    """
    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_extension == '.docx':
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {file_extension}")

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file using pypdfium2."""
    text = ""
    try:
        pdf = pdfium.PdfDocument(file_path)
        for page_number in range(len(pdf)):
            page = pdf[page_number]
            text_page = page.get_textpage()
            text += text_page.get_text() + "\n"
        return text
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        return ""

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file."""
    try:
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        return ""

def split_clauses(text: str) -> List[str]:
    """
    Split document text into individual clauses.
    
    Args:
        text: Full document text
        
    Returns:
        List of clause texts
    """
    # Simple heuristic: Split by numbered sections, articles, or clauses
    # This is a basic implementation and might need refinement based on document structure
    
    # Pattern for common clause headers like "1.", "Article 1", "Section 1", etc.
    patterns = [
        r'(?:\n|\r\n)(?:Article|Section|Clause)\s+\d+[\.\s]',
        r'(?:\n|\r\n)\d+\.(?:\d+)?[\.\s]',  # Matches "1.", "1.1.", etc.
        r'(?:\n|\r\n)(?:[A-Z])\.[\s]'  # Matches "A.", "B.", etc.
    ]
    
    # Combine patterns
    combined_pattern = '|'.join(patterns)
    
    # Split text using the combined pattern
    splits = re.split(combined_pattern, text)
    
    # Clean up splits
    clauses = [clause.strip() for clause in splits if clause.strip()]
    
    # If no clauses were detected, use paragraphs as a fallback
    if len(clauses) <= 1:
        clauses = [p.strip() for p in text.split('\n\n') if p.strip()]
    
    return clauses

def summarize_clause(clause_text: str) -> str:
    """
    Summarize a legal clause into plain English using an LLM.
    
    Args:
        clause_text: Text of the legal clause
        
    Returns:
        Simplified summary of the clause
    """
    try:
        if not clause_text or len(clause_text.strip()) < 10:
            return "No content to summarize."
            
        # If OpenAI API key is available, use it
        if openai.api_key:
            try:
                response = openai.chat.completions.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "system", "content": "You are a legal expert. Summarize the following legal clause in simple, plain English that a layperson can understand."},
                        {"role": "user", "content": clause_text}
                    ],
                    max_tokens=150
                )
                return response.choices[0].message.content.strip()
            except Exception as e:
                print(f"OpenAI API error: {e}")
                # Fall back to placeholder implementation
        
        # Placeholder implementation
        words = clause_text.split()
        if len(words) > 30:
            summary = " ".join(words[:30]) + "..."
        else:
            summary = clause_text
        return f"Summary: {summary}"
    except Exception as e:
        print(f"Error summarizing clause: {e}")
        return "Summary unavailable"

def classify_risk(clause_text: str) -> str:
    """
    Classify the risk level of a legal clause.
    
    Args:
        clause_text: Text of the legal clause
        
    Returns:
        Risk level: "High", "Medium", or "Low"
    """
    try:
        if not clause_text or len(clause_text.strip()) < 10:
            return "Medium"  # Default for empty or very short clauses
            
        # If OpenAI API key is available, use it
        if openai.api_key:
            try:
                response = openai.chat.completions.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "system", "content": "You are a legal expert. Classify the risk level of the following legal clause as 'High', 'Medium', or 'Low'."},
                        {"role": "user", "content": clause_text}
                    ],
                    max_tokens=10
                )
                risk = response.choices[0].message.content.strip()
                if risk not in ["High", "Medium", "Low"]:
                    risk = "Medium"  # Default if response is unexpected
                return risk
            except Exception as e:
                print(f"OpenAI API error in risk classification: {e}")
                # Fall back to heuristic implementation
        
        # Simple heuristic implementation
        risk_keywords = {
            "High": ["terminate", "liability", "damages", "penalty", "breach", "lawsuit", "legal action", "indemnify", 
                    "termination", "void", "violation", "prohibited", "criminal", "illegal", "fraud", "penalty"],
            "Medium": ["modify", "change", "update", "amend", "restrict", "limit", "reserve the right", 
                      "discretion", "may", "option", "choose", "elect", "decide"],
            "Low": ["contact", "notify", "inform", "provide", "notice", "communication", "information"]
        }
        
        clause_lower = clause_text.lower()
        
        # Count occurrences of risk keywords
        high_count = sum(1 for word in risk_keywords["High"] if word.lower() in clause_lower)
        medium_count = sum(1 for word in risk_keywords["Medium"] if word.lower() in clause_lower)
        low_count = sum(1 for word in risk_keywords["Low"] if word.lower() in clause_lower)
        
        # Determine risk level based on keyword counts
        if high_count > 0:
            return "High"
        elif medium_count > 0:
            return "Medium"
        elif low_count > 0:
            return "Low"
        else:
            return "Medium"  # Default to medium risk if no keywords match
    except Exception as e:
        print(f"Error classifying risk: {e}")
        return "Medium"  # Default to medium risk on error